# app/routes/markdown_converter.py
from flask import Blueprint, render_template, request, jsonify, send_file
import markdown
import io
import os
import tempfile
from docx import Document
import html2docx

markdown_bp = Blueprint('markdown', __name__)

@markdown_bp.route('/')
def index():
    """Render the markdown converter page"""
    return render_template('pages/markdown.html')

@markdown_bp.route('/convert', methods=['POST'])
def convert():
    """Convert markdown to HTML"""
    try:
        data = request.get_json()
        if not data or 'markdown' not in data:
            return jsonify({'success': False, 'error': 'No markdown provided'}), 400
        
        # Convert markdown to HTML
        md_text = data['markdown']
        html_content = markdown.markdown(
            md_text,
            extensions=['extra', 'codehilite', 'tables']
        )
        
        return jsonify({
            'success': True,
            'html': html_content
        })
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@markdown_bp.route('/download/html', methods=['POST'])
def download_html():
    """Download the converted HTML as a file"""
    try:
        md_text = request.form.get('markdown', '')
        if not md_text:
            return jsonify({'success': False, 'error': 'No markdown provided'}), 400
        
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_text,
            extensions=['extra', 'codehilite', 'tables']
        )
        
        # Wrap the HTML content in a basic HTML document
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Markdown Conversion</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; max-width: 800px; margin: 0 auto; }}
        pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        blockquote {{ border-left: 4px solid #ddd; padding-left: 15px; color: #666; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
        
        # Create a memory file-like object
        mem = io.BytesIO()
        mem.write(full_html.encode('utf-8'))
        mem.seek(0)
        
        return send_file(
            mem,
            as_attachment=True,
            download_name='markdown_conversion.html',
            mimetype='text/html'
        )
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@markdown_bp.route('/download/docx', methods=['POST'])
def download_docx():
    """Download the converted markdown as a DOCX file"""
    try:
        md_text = request.form.get('markdown', '')
        if not md_text:
            return jsonify({'success': False, 'error': 'No markdown provided'}), 400
        
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_text,
            extensions=['extra', 'codehilite', 'tables']
        )
        
        # Create a temporary HTML file
        with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as temp_html:
            temp_html_path = temp_html.name
            temp_html.write(f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Markdown Conversion</title>
</head>
<body>
{html_content}
</body>
</html>""".encode('utf-8'))
        
        # Create a temporary DOCX file
        temp_docx_path = tempfile.NamedTemporaryFile(suffix='.docx', delete=False).name
        
        try:
            # Use pandoc for conversion (must be installed on the system)
            import subprocess
            result = subprocess.run(
                ['pandoc', temp_html_path, '-f', 'html', '-t', 'docx', '-o', temp_docx_path],
                capture_output=True,
                text=True,
                check=True
            )
            
            # Return the file
            return send_file(
                temp_docx_path,
                as_attachment=True,
                download_name='markdown_conversion.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except subprocess.CalledProcessError as e:
            return jsonify({'success': False, 'error': f"Pandoc conversion failed: {e.stderr}"}), 500
        except FileNotFoundError:
            # Fallback if pandoc is not available - use python-mammoth in reverse
            from docx import Document
            
            # Create a Document
            doc = Document()
            
            # Parse HTML and convert to DOCX
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Process headings
            for i in range(1, 7):
                for heading in soup.find_all(f'h{i}'):
                    doc.add_heading(heading.get_text(), level=i)
                    heading.decompose()  # Remove the heading from the soup
            
            # Process paragraphs
            for para in soup.find_all('p'):
                p = doc.add_paragraph()
                for child in para.children:
                    if child.name == 'strong' or child.name == 'b':
                        p.add_run(child.get_text()).bold = True
                    elif child.name == 'em' or child.name == 'i':
                        p.add_run(child.get_text()).italic = True
                    else:
                        p.add_run(str(child))
            
            # Process lists
            for ul in soup.find_all('ul'):
                for li in ul.find_all('li'):
                    doc.add_paragraph(li.get_text(), style='List Bullet')
            
            for ol in soup.find_all('ol'):
                for i, li in enumerate(ol.find_all('li')):
                    doc.add_paragraph(li.get_text(), style='List Number')
            
            # Save to a BytesIO object
            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)
            
            return send_file(
                docx_io,
                as_attachment=True,
                download_name='markdown_conversion.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        finally:
            # Clean up temporary files
            try:
                os.unlink(temp_html_path)
                os.unlink(temp_docx_path)
            except:
                pass
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500